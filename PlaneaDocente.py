# PlaneaDocente - Generador de planeaciones con OpenRouter y Word

import streamlit as st
import requests
from docx import Document
import tempfile
import os

st.set_page_config(page_title="PlaneaDocente", layout="centered")
st.title("📘 PlaneaDocente con OpenRouter + Word")

# Obtener la API Key desde secrets.toml
api_key = st.secrets["OPENROUTER_API_KEY"]

# Entradas del usuario
subject = st.text_input("Asignatura (ej. Matemáticas)")
grade = st.text_input("Grado o nivel (ej. 3° primaria)")
competency = st.text_input("Competencia o aprendizaje esperado")
duration = st.text_input("Duración de clase (ej. 50 minutos)")
topic = st.text_input("Tema específico")

# Función para generar planeación con OpenRouter
@st.cache_data(show_spinner=True)
def generar_planeacion(subject, grade, competency, duration, topic):
    prompt = f"""
    Genera una planeación didáctica para una clase de {subject} en {grade}. El tema específico es \"{topic}\" y debe enfocarse en el siguiente aprendizaje esperado: \"{competency}\". La clase dura {duration}. Incluye:
    - Propósito
    - Actividades de inicio, desarrollo y cierre
    - Recursos didácticos
    - Evaluación sugerida
    Escribe en español en formato claro.
    """

    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }

        body = {
            "model": "openai/gpt-3.5-turbo",
            "messages": [{"role": "user", "content": prompt}]
        }

        response = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=body)

        if response.status_code != 200:
            return f"❌ Error {response.status_code}: {response.text}"

        data = response.json()

        if "choices" in data:
            return data["choices"][0]["message"]["content"]
        else:
            return f"❌ Respuesta inesperada: {data}"

    except Exception as e:
        return f"❌ Error al generar la planeación: {str(e)}"

# Función para crear archivo Word
def crear_docx(contenido):
    doc = Document()
    doc.add_heading("Planeación Didáctica", 0)
    for linea in contenido.split("\n"):
        if linea.strip():
            doc.add_paragraph(linea.strip())
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name

# Generar planeación
if st.button("Generar planeación"):
    if all([subject, grade, competency, duration, topic]):
        with st.spinner("Generando planeación..."):
            contenido = generar_planeacion(subject, grade, competency, duration, topic)
            if contenido.startswith("❌"):
                st.error(contenido)
            else:
                st.success("Planeación generada con éxito ✅")
                st.text_area("Vista previa:", value=contenido, height=400)

                archivo = crear_docx(contenido)
                with open(archivo, "rb") as f:
                    st.download_button("Descargar como Word (.docx)", f, file_name="planeacion.docx")
    else:
        st.warning("Por favor llena todos los campos.")
